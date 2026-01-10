"""
AI Chunkenizer - Word Document Extractor
Text extraction from .docx files using python-docx.

Copyright (c) 2025 VIEWVALUE LLC. All rights reserved.
Licensed under the MIT License. See LICENSE file for details.
"""

from docx import Document


def extract_docx(file_path: str) -> str:
    """
    Extract text from a Word document.

    Extracts both paragraph text and table content,
    preserving document structure with markers.

    Args:
        file_path: Path to the .docx file

    Returns:
        Extracted text with structure markers

    Raises:
        Exception: If document cannot be processed
    """
    try:
        doc = Document(file_path)
        text_parts = []

        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)

        # Extract tables
        for table_idx, table in enumerate(doc.tables, 1):
            table_rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                table_rows.append(" | ".join(cells))

            if table_rows:
                table_text = "\n".join(table_rows)
                text_parts.append(f"\n[Table {table_idx}]\n{table_text}\n")

        return "\n\n".join(text_parts)

    except Exception as e:
        raise Exception(f"Failed to extract Word document: {str(e)}")
